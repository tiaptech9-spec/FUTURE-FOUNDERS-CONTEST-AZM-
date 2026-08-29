from flask import Flask, render_template, request, redirect, url_for
from docx import Document
import os
import re

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

from database import (
    init_db,
    save_candidate,
    get_candidates,
    get_candidate,
    search_candidates,
    update_status,
    delete_candidate,
    clear_candidates,
)

# FLASK APPLICATION

app = Flask(__name__)


# CONFIGURATION


# This folder is used for   the uploaded resumes

UPLOAD_FOLDER = "uploads"

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

# Create uploads folder automatically if it doesn't exist

os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ALLOWED FILE

def allowed_file(filename):

    return (
        filename
        and filename.lower().endswith(".docx")
    )

# TEXT EXTRACTION FROM DOCX

def extract_text_from_docx(file_path):

    document = Document(file_path)

    paragraphs = []

    # Read normal paragraphs present in file
    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if text:
            paragraphs.append(text)

    # Read tables if it's present
    for table in document.tables:

        for row in table.rows:

            row_text = []

            for cell in row.cells:

                text = cell.text.strip()

                if text:
                    row_text.append(text)

            if row_text:

                paragraphs.append(
                    " ".join(row_text)
                )

    return "\n".join(paragraphs)

# TEXT PREPROCESSING


def preprocess_text(text):

    if not text:
        return ""

    text = text.lower()

    text = re.sub(
        r"[^a-zA-Z0-9\s]",
        " ",
        text
    )

    text = re.sub(
        r"\s+",
        " ",
        text
    )

    return text.strip()


# EMAIL EXTRACTION

def extract_email(text):

    if not text:
        return ""

    pattern = r"""
        [A-Za-z0-9._%+-]+
        @
        [A-Za-z0-9.-]+\.[A-Za-z]{2,}
    """

    match = re.search(
        pattern,
        text,
        re.VERBOSE
    )

    if match:

        return match.group(0)

    return ""


# PHONE EXTRACTION


def extract_phone(text):

    if not text:
        return ""

    pattern = r"""
        (?:\+\d{1,3}[\s-]?)?
        (?:\d{3}[\s-]?)?
        \d{3}[\s-]?\d{4}
    """

    matches = re.findall(
        pattern,
        text,
        re.VERBOSE
    )

    for match in matches:

        phone = match.strip()

        digits = re.sub(
            r"\D",
            "",
            phone
        )

        if len(digits) >= 10:

            return phone

    return ""



# CANDIDATE NAME

def extract_candidate_name(
filename,
resume_text
):

    # Use filename first
    name = os.path.splitext(filename)[0]

    name = re.sub(
        r"[_\-]+",
        " ",
        name
    )

    name = re.sub(
        r"\s+",
        " ",
        name
    ).strip()

    # Remove common resume words
    cleaned_name = re.sub(
        r"\b(resume|cv|curriculum vitae)\b",
        "",
        name,
        flags=re.IGNORECASE
    ).strip()

    if cleaned_name:

        return cleaned_name

    # Fallback to first non-empty line
    for line in resume_text.splitlines():

        line = line.strip()

        if line and len(line) <= 100:

            return line

    return "Unknown Candidate"



# TF-IDF SIMILARITY


def calculate_tfidf_score(
job_description,
resume_text
):

    job_text = preprocess_text(
        job_description
    )

    resume_text = preprocess_text(
        resume_text
    )

    if not job_text or not resume_text:

        return 0.0

    documents = [
        job_text,
        resume_text
    ]

    vectorizer = TfidfVectorizer(
        stop_words="english"
    )

    try:

        vectors = vectorizer.fit_transform(
            documents
        )

        similarity = cosine_similarity(
            vectors[0:1],
            vectors[1:2]
        )[0][0]

        return round(
            similarity * 100,
            2
        )

    except ValueError:

        return 0.0

# SEMANTIC / KEYWORD SIMILARITY


def calculate_semantic_score(
job_description,
resume_text
):

    job_words = set(
        preprocess_text(
            job_description
        ).split()
    )

    resume_words = set(
        preprocess_text(
            resume_text
        ).split()
    )

    if not job_words or not resume_words:

        return 0.0

    common_words = job_words.intersection(
        resume_words
    )

    score = (
        len(common_words)
        / len(job_words)
    ) * 100

    return round(
        min(score, 100),
        2
    )


# FINAL SCORE


def calculate_final_score(
tfidf_score,
semantic_score
):

    final_score = (
        (tfidf_score * 0.5)
        +
        (semantic_score * 0.5)
    )

    return round(
        final_score,
        2
    )



# CANDIDATE STATUS


def get_candidate_status(score):

    if score >= 70:

        return "Shortlisted"

    elif score >= 50:

        return "Under Review"

    return "Rejected"

# PREPARE CANDIDATE DATA WITH SCORES, RANKING, STATUS, AND RESUME DETAILS.


def prepare_candidate(
candidate,
rank=None
):

    if not candidate:

        return None

    candidate = dict(candidate)

    if rank is None:

        rank = candidate.get(
            "id",
            0
        )

    overall_score = float(
        candidate.get(
            "overall_score",
            0
        ) or 0
    )

    return {

        "id": candidate.get("id"),

        "name": (
            candidate.get("candidate_name")
            or "Unknown Candidate"
        ),

        "filename": (
            candidate.get("filename")
            or ""
        ),

        "email": (
            candidate.get("email")
            or ""
        ),

        "phone": (
            candidate.get("phone")
            or ""
        ),

        "resume_text": (
            candidate.get("resume_text")
            or ""
        ),

        "clean_text": (
            candidate.get("clean_text")
            or ""
        ),

        "tfidf_score": round(
            float(
                candidate.get(
                    "tfidf_score",
                    0
                ) or 0
            ),
            2
        ),

        "semantic_score": round(
            float(
                candidate.get(
                    "semantic_score",
                    0
                ) or 0
            ),
            2
        ),

        "final_score": round(
            overall_score,
            2
        ),

        "status": (
            candidate.get("status")
            or get_candidate_status(
                overall_score
            )
        ),

        "uploaded_at": (
            candidate.get("uploaded_at")
            or ""
        ),

        "rank": rank
    }


# HOME / SCREENING PAGE

@app.route(
"/",
methods=["GET", "POST"]
)
def index():

    if request.method == "POST":

        job_description = request.form.get(
            "job_description",
            ""
        ).strip()

        uploaded_files = request.files.getlist(
            "resumes"
        )

        
        # CHECK JOB DESCRIPTION
   

        if not job_description:

            return render_template(
                "index.html",
                error="Please enter a job description."
            )

        # CHECK FILES
       

        valid_files = [

            file

            for file in uploaded_files

            if file
            and file.filename
            and allowed_file(file.filename)

        ]

        if not valid_files:

            return render_template(
                "index.html",
                error="Please upload at least one DOCX resume."
            )

        
        # PROCESS EACH RESUME

        for file in valid_files:

            filename = file.filename

            # Keep only safe filename characters
            safe_filename = re.sub(
                r"[^a-zA-Z0-9._\- ]",
                "_",
                filename
            )

            # IMPORTANT
            # Resume is saved ONLY inside uploads folder
            file_path = os.path.join(
                app.config["UPLOAD_FOLDER"],
                safe_filename
            )

            file.save(file_path)

            try:

                # EXTRACT RESUME TEXT

                resume_text = extract_text_from_docx(
                    file_path
                )

               
                # PREPROCESS TEXT
                

                clean_text = preprocess_text(
                    resume_text
                )

                
                # EXTRACT CANDIDATE INFORMATION
                

                candidate_name = extract_candidate_name(
                    safe_filename,
                    resume_text
                )

                email = extract_email(
                    resume_text
                )

                phone = extract_phone(
                    resume_text
                )

                
                # CALCULATE ML SCORES
                

                tfidf_score = calculate_tfidf_score(
                    job_description,
                    resume_text
                )

                semantic_score = calculate_semantic_score(
                    job_description,
                    resume_text
                )

                overall_score = calculate_final_score(
                    tfidf_score,
                    semantic_score
                )

                # SAVE RESULT TO DATABASE
                

                save_candidate(
                    filename=safe_filename,
                    candidate_name=candidate_name,
                    email=email,
                    phone=phone,
                    resume_text=resume_text,
                    clean_text=clean_text,
                    tfidf_score=tfidf_score,
                    semantic_score=semantic_score,
                    overall_score=overall_score
                )

            except Exception as error:

                print(
                    f"Error processing {filename}: {error}"
                )

        return redirect(
            url_for("dashboard")
        )

    return render_template(
        "index.html"
    )


# DASHBOARD



@app.route("/dashboard")
def dashboard():

    search = request.args.get(
        "search",
        ""
    ).strip()

    status = request.args.get(
        "status",
        ""
    ).strip()

    min_score_text = request.args.get(
        "min_score",
        "0"
    ).strip()

    try:

        min_score = float(
            min_score_text
        )

    except ValueError:

        min_score = 0

    rows = search_candidates(
        search=search,
        min_score=min_score,
        status=status
    )

    candidates = []

    for index, row in enumerate(
        rows,
        start=1
    ):

        candidates.append(
            prepare_candidate(
                row,
                rank=index
            )
        )

    
    # DASHBOARD STATISTICS
    

    all_rows = get_candidates()

    total_candidates = len(
        all_rows
    )

    shortlisted = sum(
        1
        for row in all_rows
        if row["status"] == "Shortlisted"
    )

    under_review = sum(
        1
        for row in all_rows
        if row["status"] == "Under Review"
    )

    rejected = sum(
        1
        for row in all_rows
        if row["status"] == "Rejected"
    )

    if all_rows:

        average_score = round(

            sum(
                float(
                    row["overall_score"] or 0
                )
                for row in all_rows
            )
            /
            len(all_rows),

            2
        )

    else:

        average_score = 0

    return render_template(
        "dashboard.html",
        candidates=candidates,
        total_candidates=total_candidates,
        shortlisted=shortlisted,
        under_review=under_review,
        rejected=rejected,
        average_score=average_score,
        search=search,
        status=status,
        min_score=min_score
    )


# CANDIDATE DETAILS


@app.route(
"/candidate/<int:candidate_id>"
)
def candidate(candidate_id):

    selected_candidate = get_candidate(
        candidate_id
    )

    if selected_candidate:

        all_rows = get_candidates()

        selected_score = float(
            selected_candidate["overall_score"]
            or 0
        )

        rank = 1

        for row in all_rows:

            row_score = float(
                row["overall_score"]
                or 0
            )

            if row_score > selected_score:

                rank += 1

        selected_candidate = prepare_candidate(
            selected_candidate,
            rank=rank
        )

    return render_template(
        "candidate.html",
        candidate=selected_candidate
    )



# UPDATE CANDIDATE STATUS


@app.route(
"/candidate/<int:candidate_id>/status",
methods=["POST"]
)
def change_status(candidate_id):

    status = request.form.get(
        "status",
        ""
    ).strip()

    allowed_statuses = [
        "Shortlisted",
        "Under Review",
        "Rejected"
    ]

    if status in allowed_statuses:

        update_status(
            candidate_id,
            status
        )

    return redirect(
        url_for(
            "candidate",
            candidate_id=candidate_id
        )
    )



# DELETE CANDIDATE


@app.route(
"/candidate/<int:candidate_id>/delete",
methods=["POST"]
)
def remove_candidate(candidate_id):

    candidate = get_candidate(
        candidate_id
    )

    if candidate:

        filename = candidate["filename"]

        delete_candidate(
            candidate_id
        )

        # Delete uploaded resume from uploads/
        file_path = os.path.join(
            app.config["UPLOAD_FOLDER"],
            filename
        )

        if os.path.exists(file_path):

            try:

                os.remove(file_path)

            except OSError:

                pass

    return redirect(
        url_for("dashboard")
    )


# CLEAR ALL CANDIDATES



@app.route(
"/clear-candidates",
methods=["POST"]
)
def clear_all_candidates():

    clear_candidates()

    # Delete all uploaded resumes from uploads/
    for filename in os.listdir(
        app.config["UPLOAD_FOLDER"]
    ):

        file_path = os.path.join(
            app.config["UPLOAD_FOLDER"],
            filename
        )

        if os.path.isfile(file_path):

            try:

                os.remove(file_path)

            except OSError:

                pass

    return redirect(
        url_for("dashboard")
    )


# RUN APPLICATION
if __name__ == "__main__":

    # Automatically create database and a table
    init_db()

    app.run(
        debug=True
    )
